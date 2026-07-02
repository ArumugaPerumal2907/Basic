"""
document_parser.py
-------------------
Extracts raw text from uploaded documents so the rest of the pipeline
(detectors, risk classifier, AI service) can work with a single
plain-text representation regardless of source format.

Supported formats: PDF, TXT, MD, CSV, DOCX, XLSX/XLS, JSON, PNG/JPG/JPEG
(the last three via the OCR layer in ocr_service.py).

Scanned (image-only) PDFs are auto-detected via a simple heuristic --
very little extractable text relative to page count -- and routed
through OCR automatically when a Sarvam Vision key is available.
"""

import io
import json
from typing import Optional

import pandas as pd
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = [
    "pdf", "txt", "md", "csv", "docx", "xlsx", "xls", "json",
    "png", "jpg", "jpeg",
]


def parse_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def parse_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages.append(f"\n--- Page {i + 1} ---\n{text}")
    return "\n".join(pages)


def parse_csv(file_bytes: bytes) -> str:
    df = pd.read_csv(io.BytesIO(file_bytes))
    # Represent the CSV both as a readable table (for the AI summary /
    # display) and as raw cell text (so every value is still scanned by
    # the regex detectors, including values pandas may reformat).
    table_repr = df.to_string(index=False)
    raw_cells = "\n".join(
        " ".join(str(v) for v in row) for row in df.itertuples(index=False)
    )
    return f"{table_repr}\n\n--- Raw Cell Values ---\n{raw_cells}"


def parse_xlsx(file_bytes: bytes) -> str:
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    parts = []
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        parts.append(f"--- Sheet: {sheet} ---")
        parts.append(df.to_string(index=False))
        raw_cells = "\n".join(
            " ".join(str(v) for v in row) for row in df.itertuples(index=False)
        )
        parts.append(f"--- Sheet: {sheet} (Raw Cell Values) ---\n{raw_cells}")
    return "\n\n".join(parts)


def parse_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def parse_json(file_bytes: bytes) -> str:
    data = json.loads(file_bytes.decode("utf-8", errors="ignore"))
    # Pretty-printed so keys/values are still individually visible to the
    # regex detectors (rather than one long unbroken line).
    return json.dumps(data, indent=2, ensure_ascii=False)


def parse_image(
    file_bytes: bytes, filename: str, sarvam_api_key: Optional[str] = None
) -> str:
    from ocr_service import ocr_extract
    return ocr_extract(file_bytes, filename, sarvam_api_key=sarvam_api_key)


def _looks_scanned(text: str, min_chars_per_page: int = 20) -> bool:
    """Heuristic: a PDF yielding very little extractable text relative to
    its page count is very likely a scanned/image-based PDF rather than a
    text PDF, and should be routed through OCR instead."""
    page_count = max(text.count("--- Page"), 1)
    return len(text.strip()) < min_chars_per_page * page_count


def parse_document(
    filename: str, file_bytes: bytes, sarvam_api_key: Optional[str] = None
) -> str:
    lower = filename.lower()

    if lower.endswith(".pdf"):
        text = parse_pdf(file_bytes)
        if _looks_scanned(text) and sarvam_api_key:
            try:
                from ocr_service import ocr_extract
                ocr_text = ocr_extract(file_bytes, filename, sarvam_api_key=sarvam_api_key)
                if ocr_text.strip():
                    return ocr_text
            except Exception:
                # Fall back to whatever (possibly empty) text pypdf found
                # rather than failing the whole upload.
                pass
        return text
    elif lower.endswith(".csv"):
        return parse_csv(file_bytes)
    elif lower.endswith((".txt", ".md")):
        return parse_txt(file_bytes)
    elif lower.endswith(".docx"):
        return parse_docx(file_bytes)
    elif lower.endswith((".xlsx", ".xls")):
        return parse_xlsx(file_bytes)
    elif lower.endswith(".json"):
        return parse_json(file_bytes)
    elif lower.endswith((".png", ".jpg", ".jpeg")):
        return parse_image(file_bytes, filename, sarvam_api_key=sarvam_api_key)
    else:
        raise ValueError(f"Unsupported file type: {filename}")
